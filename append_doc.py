from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(r"TASKS\Phase 2\week 1_day 1\Assignment.docx")

# Add spacing before new content
doc.add_paragraph()

# Title
title = doc.add_heading("Phase 2 – Week 1, Day 1: Building an AI Flow in Langflow", level=1)

# Intro
doc.add_paragraph(
    "This assignment demonstrates building a simple AI-powered document processing "
    "pipeline using Langflow (a visual, no-code/low-code tool for chaining AI components)."
)

# Section: What the Flow Does
doc.add_heading("What the Flow Does", level=2)
doc.add_paragraph(
    "The pipeline reads a PDF file, sends its content to an LLM (via Groq), and outputs "
    "the result. Here's how each component works:"
)

# Image 1
doc.add_heading("Image 1 — The Flow Canvas:", level=3)
doc.add_paragraph("The flow has three connected nodes:")

# Node 1
p1 = doc.add_paragraph()
run1_num = p1.add_run("1. ")
run1_bold = p1.add_run("Read File")
run1_bold.bold = True
p1.add_run(
    " — Loads a PDF (sample_invoice (1).pdf, 3.65 KB) and extracts its raw text content. "
    'The "Advanced Parser" option is toggled off, meaning it uses basic text extraction.'
)

# Node 2
p2 = doc.add_paragraph()
run2_num = p2.add_run("2. ")
run2_bold = p2.add_run("Prompt Template")
run2_bold.bold = True
p2.add_run(
    " — Takes the raw file content and injects it into a prompt. The visible text suggests "
    'a prompt like "Read the entire PDF get strictly Copywriting – Homepage and..." — '
    "likely asking the LLM to extract or summarize specific information from the invoice."
)

# Node 3
p3 = doc.add_paragraph()
run3_num = p3.add_run("3. ")
run3_bold = p3.add_run("Groq (LLM)")
run3_bold.bold = True
p3.add_run(
    " — Sends the assembled prompt to the Llama-3.3-70b-versatile model via Groq's API. "
    "Key settings visible on the right panel:"
)

# Sub-bullets for Groq
doc.add_paragraph("    • Temperature: ~0.19 (low = more factual/deterministic output)")
doc.add_paragraph("    • Groq API Key is connected")
doc.add_paragraph("    • Output feeds into a Chat Output node")

# Image 2
doc.add_heading("Image 2 — The Result:", level=3)
doc.add_paragraph(
    "After running the flow, the Component Output panel shows the result: $600.00 — "
    "meaning the AI successfully read the invoice PDF and extracted the total amount ($600) from it."
)

doc.save(r"TASKS\Phase 2\week 1_day 1\Assignment.docx")
print("Done - text appended to Assignment.docx")
